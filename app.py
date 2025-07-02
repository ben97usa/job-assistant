from flask import Flask, render_template, request, session, redirect, url_for
import os
import openai
from dotenv import load_dotenv
import fitz  # PyMuPDF
import docx
from docx import Document
import PyPDF2

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = "uploads"
load_dotenv()
app.secret_key = "supersecretkey_ngoctoan_123456"


client = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

def extract_text(file_path, filename):
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        text = ""
        doc = fitz.open(file_path)
        for page in doc:
            text += page.get_text()
        return text
    elif ext == "docx":
        doc = docx.Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs])
    else:
        return ""

@app.route("/", methods=["GET"])
def index():
    return render_template(
        "index.html",
        name=session.get("name", ""),
        email=session.get("email", ""),
        summary=session.get("summary", ""),
        jobdesc=session.get("jobdesc", "")
    )

@app.route("/generate", methods=["POST"])
def generate():
    action = request.form.get("action")
    name = request.form["name"]
    email = request.form["email"]
    summary = request.form["summary"]
    jobdesc = request.form["jobdesc"]

    # 💾 Lưu vào session
    session["name"] = name
    session["email"] = email
    session["summary"] = summary
    session["jobdesc"] = jobdesc

    resume_text = ""
    if "resume" in request.files:
        resume = request.files["resume"]
        if resume and resume.filename != "":
            path = os.path.join(app.config['UPLOAD_FOLDER'], resume.filename)
            resume.save(path)
            resume_text = extract_text(path, resume.filename)

    try:
        if action == "optimize":
            prompt = (
                f"My name is {name}, email: {email}.\n\n"
                f"Summary:\n{summary}\n\n"
                f"Job Description:\n{jobdesc}\n\n"
                f"My current resume:\n{resume_text}\n\n"
                f"Can you rewrite and optimize my resume for ATS based on this job description?"
            )
        elif action == "learn":
            prompt = (
                f"I am {name}, my email is {email}.\n"
                f"I want to learn how to write an effective resume from scratch based on this job description:\n{jobdesc}\n\n"
                f"Please guide me step by step with examples."
            )
        else:
            return render_template("result.html", cover_letter="Invalid action.")

        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}]
        )
        result = response.choices[0].message.content
    except Exception as e:
        result = f"Error: {str(e)}"

    return render_template("result.html", cover_letter=result)

@app.route("/learn", methods=["POST"])
def learn():
    try:
        resume_file = request.files["resume"]
        filename = resume_file.filename.lower()

        if filename.endswith(".docx"):
            doc = Document(resume_file)
            resume_text = "\n".join([p.text for p in doc.paragraphs])
        elif filename.endswith(".pdf"):
            reader = PyPDF2.PdfReader(resume_file)
            resume_text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
        else:
            resume_text = "Unsupported file format. Please upload .docx or .pdf only."

        prompt = f"Improve this resume to match modern ATS standards for better interview chances:\n\n{resume_text}"
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}]
        )
        result = response.choices[0].message.content

    except Exception as e:
        result = f"Error: {str(e)}"

    return render_template("result.html", cover_letter=result)


if __name__ == "__main__":
    app.run(debug=True)
